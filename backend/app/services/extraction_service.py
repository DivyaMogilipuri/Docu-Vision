import fitz
import pytesseract
from pytesseract import Output
from pdf2image import convert_from_path
from docx import Document
from PIL import Image
import uuid
import re
import io
import numpy as np
import cv2
from typing import List, Dict

POPPLER_PATH = r"C:/Program Files/poppler-25.11.0/Library/bin"

TESS_CONFIG = "--oem 3 --psm 6 -l eng"

import fitz
import pytesseract
from pytesseract import Output
from PIL import Image
import io
import uuid
import cv2
import numpy as np

import fitz
import pytesseract
from pytesseract import Output
from docx import Document
from PIL import Image
import uuid
import re
import io
from typing import List, Dict

# =====================================================
# DIGITAL TEXT EXTRACTION
# =====================================================

def extract_digital_blocks(page, page_number, document_id):

    blocks = []
    text_dict = page.get_text("dict")

    for block in text_dict["blocks"]:

        if block["type"] == 0:

            for line in block["lines"]:

                line_text = ""
                x0 = y0 = x1 = y1 = None

                for span in line["spans"]:

                    if not line_text:
                        x0, y0 = span["bbox"][0], span["bbox"][1]

                    x1, y1 = span["bbox"][2], span["bbox"][3]

                    line_text += span["text"] + " "

                print("page number from digital blocks",page_number)

                if line_text.strip():

                    blocks.append({
                        "block_id": str(uuid.uuid4()),
                        "document_id": document_id,
                        "type": "digital_text",
                        "page_number": page_number,
                        "text": line_text.strip(),
                        "x0": x0,
                        "y0": y0,
                        "x1": x1,
                        "y1": y1
                    })

    return blocks


# =====================================================
# OCR WORD EXTRACTION
# =====================================================

def extract_words_from_image(image, page, page_number, document_id):

    blocks = []

    ocr_data = pytesseract.image_to_data(image, output_type=Output.DICT)

    scale_x = page.rect.width / image.width
    scale_y = page.rect.height / image.height

    for i in range(len(ocr_data["text"])):

        word = ocr_data["text"][i].strip()

        try:
            conf = float(ocr_data["conf"][i])
        except:
            conf = -1

        if word and conf > 20:

            x = ocr_data["left"][i]
            y = ocr_data["top"][i]
            w = ocr_data["width"][i]
            h = ocr_data["height"][i]

            print("page no. from extract words frm image",page_number)

            blocks.append({
                "block_id": str(uuid.uuid4()),
                "document_id": document_id,
                "type": "ocr_word",
                "page_number": page_number,
                "text": word,
                "x0": x * scale_x,
                "y0": y * scale_y,
                "x1": (x + w) * scale_x,
                "y1": (y + h) * scale_y
            })

    return blocks


# =====================================================
# GENERIC TEXT REGION DETECTION
# =====================================================

def detect_text_regions(word_blocks, x_threshold=150, y_threshold=50):

    regions = []

    for block in word_blocks:

        bx0, by0, bx1, by1 = block["x0"], block["y0"], block["x1"], block["y1"]

        placed = False

        for region in regions:

            rx0, ry0, rx1, ry1 = region

            if abs(bx0 - rx0) < x_threshold and abs(by0 - ry0) < y_threshold:

                region[0] = min(rx0, bx0)
                region[1] = min(ry0, by0)
                region[2] = max(rx1, bx1)
                region[3] = max(ry1, by1)

                placed = True
                break

        if not placed:
            regions.append([bx0, by0, bx1, by1])

    return regions


# =====================================================
# CROP REGIONS
# =====================================================

def crop_regions(page, regions):

    cropped_images = []

    for r in regions:

        rect = fitz.Rect(r[0], r[1], r[2], r[3])

        pix = page.get_pixmap(clip=rect, dpi=300)

        image = Image.open(io.BytesIO(pix.tobytes()))

        cropped_images.append(image)

    return cropped_images


# =====================================================
# GROUP WORDS INTO LINES
# =====================================================

def group_words_into_lines(blocks, y_threshold=12):

    if not blocks:
        return []

    blocks = sorted(
        blocks,
        key=lambda b: (
            b["page_number"],
            (b["y0"] + b["y1"]) / 2,
            b["x0"]
        )
    )

    lines = []
    current_line = []
    current_y = None

    for block in blocks:

        center_y = (block["y0"] + block["y1"]) / 2

        if current_y is None:
            current_y = center_y
            current_line.append(block)

        elif abs(center_y - current_y) <= y_threshold:
            current_line.append(block)

        else:
            lines.append(current_line)
            current_line = [block]
            current_y = center_y

    if current_line:
        lines.append(current_line)

    merged_lines = []

    for line in lines:

        line = sorted(line, key=lambda b: b["x0"])

        text = " ".join(b["text"] for b in line)

        print("page nuber from group words into lines ",line[0]["page_number"])

        merged_lines.append({
            "block_id": str(uuid.uuid4()),
            "document_id": line[0]["document_id"],
            "type": "ocr_text",
            "page_number": line[0]["page_number"],
            "text": text,
            "x0": min(b["x0"] for b in line),
            "y0": min(b["y0"] for b in line),
            "x1": max(b["x1"] for b in line),
            "y1": max(b["y1"] for b in line)
        })

    return merged_lines
# =====================================================
# DETECT COLUMNS
# =====================================================

def detect_columns(page_blocks, page_width):

    if not page_blocks:
        return [page_blocks]

    # sort blocks left → right
    blocks_sorted = sorted(page_blocks, key=lambda b: b["x0"])

    # dynamic column gap threshold
    column_gap = page_width * 0.15

    columns = []
    current_column = [blocks_sorted[0]]

    prev_x = blocks_sorted[0]["x0"]

    for block in blocks_sorted[1:]:
        gap = block["x0"] - prev_x

        if gap > column_gap:
            columns.append(current_column)
            current_column = [block]
        else:
            current_column.append(block)

        prev_x = block["x0"]

    columns.append(current_column)

    return columns

def merge_lines_into_paragraphs(blocks, y_gap_threshold=18, x_align_threshold=30):

    if not blocks:
        return []

    merged_blocks = []
    current_para = blocks[0].copy()

    for block in blocks[1:]:

        text = block["text"].strip()

        vertical_gap = block["y0"] - current_para["y1"]

        x_aligned = abs(block["x0"] - current_para["x0"]) < x_align_threshold

        prev_text = current_para["text"].strip()

        end_punctuation = prev_text.endswith((".", "!", "?", ":"))

        is_bullet = bool(re.match(r"^[-•●*]", text))

        if (
            vertical_gap < y_gap_threshold
            and x_aligned
            and not end_punctuation
            and not is_bullet
        ):

            current_para["text"] += " " + text

            current_para["x0"] = min(current_para["x0"], block["x0"])
            current_para["y0"] = min(current_para["y0"], block["y0"])
            current_para["x1"] = max(current_para["x1"], block["x1"])
            current_para["y1"] = max(current_para["y1"], block["y1"])

        else:

            merged_blocks.append(current_para)
            current_para = block.copy()

    merged_blocks.append(current_para)

    return merged_blocks
# =====================================================
# LAYOUT SORT
# =====================================================
def layout_sort(all_blocks, doc):

    final_blocks = []
    pages = {}

    for b in all_blocks:
        pages.setdefault(b["page_number"], []).append(b)

    for page_number in sorted(pages.keys()):

        page_blocks = pages[page_number]

        page = doc[page_number - 1]
        page_width = page.rect.width

        columns = detect_columns(page_blocks, page_width)

        # sort columns left → right
        columns = sorted(columns, key=lambda col: min(b["x0"] for b in col))

        for col in columns:

            # sort top → bottom
            col_sorted = sorted(col, key=lambda b: (b["y0"], b["x0"]))

            # MERGE PARAGRAPHS INSIDE COLUMN
            merged = merge_lines_into_paragraphs(col_sorted)

            final_blocks.extend(merged)

    return final_blocks
# def layout_sort(all_blocks, doc):

#     final_blocks = []
#     pages = {}

#     for b in all_blocks:
#         pages.setdefault(b["page_number"], []).append(b)

#     for page_number in sorted(pages.keys()):
#         print("page number from layout sort",page_number)

#         page_blocks = pages[page_number]

#         page = doc[page_number - 1]
#         page_width = page.rect.width

#         columns = detect_columns(page_blocks, page_width)

#         # sort columns left → right
#         columns = sorted(columns, key=lambda col: min(b["x0"] for b in col))

#         for col in columns:
#             col_sorted = sorted(col, key=lambda b: (b["y0"], b["x0"]))
#             final_blocks.extend(col_sorted)

#     return final_blocks



# =====================================================
# MAIN PDF PIPELINE
# =====================================================

def extract_from_pdf_layout_aware(pdf_path):

    document_id = str(uuid.uuid4())

    doc = fitz.open(pdf_path)

    all_blocks = []

    for page_index, page in enumerate(doc):

        page_number = page_index + 1
        
        print("page number from extract pdf aware",page_number,"page index ",page_index)

        digital_blocks = extract_digital_blocks(page, page_number, document_id)

        if digital_blocks:

            all_blocks.extend(digital_blocks)
            continue

        # If no digital text → run OCR pipeline

        pix = page.get_pixmap(dpi=300)

        page_image = Image.open(io.BytesIO(pix.tobytes()))


        word_blocks = extract_words_from_image(page_image, page, page_number, document_id)

        regions = detect_text_regions(word_blocks)

        cropped_images = crop_regions(page, regions)

        region_words = []

        for img in cropped_images:

            region_words.extend(
                extract_words_from_image(img, page, page_number, document_id)
            )

        line_blocks = group_words_into_lines(region_words)

        all_blocks.extend(line_blocks)

    all_blocks = layout_sort(all_blocks, doc)
    
    for b in all_blocks[:]:
        print("from all PAGE:", b.get("page_number"), "| TEXT:", b.get("text"))

    return {
        "document_id": document_id,
        "doc_type": "general",
        "source": "pdf",
        "blocks": all_blocks
    }


# =====================================================
# DOCX EXTRACTION
# =====================================================

def extract_from_docx(filepath: str, doc_type="general"):

    document_id = str(uuid.uuid4())

    document = Document(filepath)

    blocks = []

    for para in document.paragraphs:

        text = re.sub(r"\s+", " ", para.text).strip()

        if text:

            blocks.append({
                "block_id": str(uuid.uuid4()),
                "document_id": document_id,
                "type": "paragraph",
                "page_number": None,
                "text": text
            })

    for table in document.tables:

        table_rows = []

        for row in table.rows:

            row_data = [
                re.sub(r"\s+", " ", cell.text).strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if row_data:
                table_rows.append(" | ".join(row_data))

        if table_rows:

            blocks.append({
                "block_id": str(uuid.uuid4()),
                "document_id": document_id,
                "type": "table",
                "page_number": None,
                "text": "\n".join(table_rows)
            })

    return {
        "document_id": document_id,
        "doc_type": doc_type,
        "source": "docx",
        "blocks": blocks
    }